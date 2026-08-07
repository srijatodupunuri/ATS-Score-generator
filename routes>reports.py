from io import BytesIO

import pandas as pd  # noqa: F401
from docx import Document
from flask import Blueprint, jsonify, render_template, send_file

from routes.report_store import load_report_history, load_reports

reports = Blueprint(
    "reports",
    __name__
)


@reports.route("/reports")
def reports_page():

    return render_template(
        "reports.html"
    )


@reports.route("/api/reports")
def get_reports():

    reports_data = load_reports()

    accepted = [
        report
        for report in reports_data
        if report["ats_result"]
        ["decision"]["status"]
        == "Accepted"
    ]

    rejected = [
        report
        for report in reports_data
        if report["ats_result"]
        ["decision"]["status"]
        == "Rejected"
    ]

    review = [
        report
        for report in reports_data
        if report["ats_result"]
        ["decision"]["status"]
        == "Review Required"
    ]

    return jsonify({

        "success": True,

        "summary": {

            "total":
                len(reports_data),

            "accepted":
                len(accepted),

            "rejected":
                len(rejected),

            "review":
                len(review)
        },

        "accepted_reports":
            accepted,

        "rejected_reports":
            rejected,

        "review_reports":
            review,

        "history":
            reports_data
    })
@reports.route(
    "/api/report-history"
)
def report_history():

    return jsonify({

        "success": True,

        "data":
        load_report_history()

    })
@reports.route(
    "/api/download/accepted"
)
def download_accepted():

    reports_data = load_reports()

    accepted = [

        candidate

        for candidate in reports_data

        if candidate["ats_result"]
            ["decision"]["status"]

        == "Accepted"

    ]

    file_buffer = generate_docx_report(

        accepted,

        "Accepted Candidates Report"
    )

    return send_file(

        file_buffer,

        as_attachment=True,

        download_name=
        "Accepted_Candidates_Report.docx",

        mimetype=
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    )
@reports.route(
    "/api/download/rejected"
)
def download_rejected():

    reports_data = load_reports()

    rejected = [

        candidate

        for candidate in reports_data

        if candidate["ats_result"]
            ["decision"]["status"]

        == "Rejected"

    ]

    file_buffer = generate_docx_report(

        rejected,

        "Rejected Candidates Report"
    )

    return send_file(

        file_buffer,

        as_attachment=True,

        download_name=
        "Rejected_Candidates_Report.docx",

        mimetype=
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    )

@reports.route(
    "/api/download/review"
)
def download_review():

    reports_data = load_reports()

    review = [

        candidate

        for candidate in reports_data

        if candidate["ats_result"]
            ["decision"]["status"]

        == "Review Required"

    ]

    file_buffer = generate_docx_report(

        review,

        "Review Required Candidates Report"
    )

    return send_file(

        file_buffer,

        as_attachment=True,

        download_name=
        "Review_Required_Candidates_Report.docx",

        mimetype=
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    )

def generate_docx_report(
    candidates,
    report_title
):

    document = Document()

    document.add_heading(
        report_title,
        level=1
    )

    for candidate in candidates:

        profile = candidate[
            "candidate_profile"
        ]

        ats = candidate[
            "ats_result"
        ]

        document.add_heading(
            profile["name"],
            level=2
        )

        document.add_paragraph(
            f"Rank: {candidate['rank']}"
        )

        document.add_paragraph(
            f"ATS Score: {ats['ats_percentage']}%"
        )

        document.add_paragraph(
            f"Status: {ats['decision']['status']}"
        )

        document.add_paragraph(
            f"Priority: {ats['decision']['priority']}"
        )

        document.add_paragraph(
            f"Reason: {ats['decision']['reason']}"
        )

        document.add_paragraph(
            f"Experience: "
            f"{profile['experience']['years']} Years"
        )

        document.add_paragraph(
            "Skills: " +
            ", ".join(
                profile["skills"]
            )
        )

        document.add_paragraph(
            "Matched Skills: " +
            ", ".join(
                ats["matched_skills"]
            )
        )

        document.add_paragraph(
            "Missing Skills: " +
            (
                ", ".join(
                    ats["missing_skills"]
                )
                if ats["missing_skills"]
                else "None"
            )
        )

        document.add_page_break()

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer
